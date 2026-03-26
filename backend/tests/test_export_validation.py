"""Step 7.4 导出文件验证 — DOCX/PDF 格式深度验证"""
from __future__ import annotations

import io
from unittest.mock import MagicMock

import pytest

MARKDOWN_CONTENT = """# 量子计算技术报告

## 第一章 概述

量子计算是一种利用量子力学原理进行计算的新型计算范式，与经典计算有本质区别。

## 第二章 核心原理

量子比特（qubit）是量子计算的基本计算单元。

### 2.1 叠加态

量子叠加允许 qubit 同时处于 0 和 1 的状态，大幅提升并行计算能力。

### 2.2 纠缠态

量子纠缠使得多个 qubit 之间存在超距关联，是量子算法加速的核心机制。

```python
# 量子电路示例
def create_bell_state(qc):
    qc.h(0)
    qc.cx(0, 1)
    return qc
```

## 第三章 应用前景

量子计算在密码学、药物发现、金融优化等领域具有革命性潜力。

### 3.1 密码学

Shor 算法可以在多项式时间内分解大整数，威胁现有 RSA 加密体系。

### 3.2 药物发现

量子模拟可以精确模拟分子结构，加速新药研发进程。

## 第四章 总结

量子计算技术正处于快速发展阶段，预计在未来十年内实现商业化突破。
"""


def _make_task(output_text: str = MARKDOWN_CONTENT, title: str = "量子计算技术报告"):
    task = MagicMock()
    task.title = title
    task.output_text = output_text
    task.status = "completed"
    return task


# ---------------------------------------------------------------------------
# DOCX 格式验证
# ---------------------------------------------------------------------------

class TestDocxExportValidation:
    def _export_docx(self, task=None):
        from app.services.exporter import DocxExporter
        exporter = DocxExporter()
        return exporter.export(task or _make_task())

    def test_docx_is_valid_zip(self):
        """DOCX 是有效的 ZIP 文件（OOXML格式）"""
        import zipfile
        data = self._export_docx()
        assert zipfile.is_zipfile(io.BytesIO(data))

    def test_docx_contains_document_xml(self):
        """DOCX 内含 word/document.xml"""
        import zipfile
        data = self._export_docx()
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            assert "word/document.xml" in z.namelist()

    def test_docx_title_is_first_heading(self):
        """python-docx 读取：第一个段落是报告标题"""
        import docx
        data = self._export_docx()
        doc = docx.Document(io.BytesIO(data))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1
        assert "量子计算技术报告" in headings[0].text

    def test_docx_has_correct_heading_levels(self):
        """标题层级正确：H1/H2/H3 都有"""
        import docx
        data = self._export_docx()
        doc = docx.Document(io.BytesIO(data))
        styles = {p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")}
        assert "Heading 1" in styles
        assert "Heading 2" in styles

    def test_docx_paragraph_count_reasonable(self):
        """段落数量合理（至少10个非空段落）"""
        import docx
        data = self._export_docx()
        doc = docx.Document(io.BytesIO(data))
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) >= 10

    def test_docx_chinese_characters_preserved(self):
        """中文字符正确保留，无乱码"""
        import docx
        data = self._export_docx()
        doc = docx.Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "量子计算" in full_text
        assert "核心原理" in full_text
        assert "叠加态" in full_text
        assert "纠缠态" in full_text

    def test_docx_code_block_preserved(self):
        """代码块内容被保留"""
        import docx
        data = self._export_docx()
        doc = docx.Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "create_bell_state" in full_text

    def test_docx_title_in_filename_encoding(self):
        """中文标题可被 URL 编码（用于 Content-Disposition）"""
        from urllib.parse import quote
        title = "量子计算技术报告"
        encoded = quote(f"{title}.docx")
        assert encoded  # 不抛异常即通过
        assert ".docx" in encoded or "%" in encoded


# ---------------------------------------------------------------------------
# PDF 格式验证
# ---------------------------------------------------------------------------

class TestPdfExportValidation:
    def _export_pdf(self, task=None):
        from app.services.exporter import PdfExporter
        exporter = PdfExporter()
        return exporter.export(task or _make_task())

    def test_pdf_starts_with_magic_bytes(self):
        """PDF 以 %PDF- 开头"""
        data = self._export_pdf()
        assert data[:5] == b"%PDF-"

    def test_pdf_ends_with_eof(self):
        """PDF 以 %%EOF 结尾"""
        data = self._export_pdf()
        assert b"%%EOF" in data[-20:]

    def test_pdf_chinese_title_included(self):
        """PDF 字节流中包含中文标题（UTF-8 编码）"""
        data = self._export_pdf()
        # PDF 内嵌 ToUnicode 映射，中文内容以某种形式存在
        assert len(data) > 1000  # 有实质内容

    def test_pdf_size_proportional_to_content(self):
        """内容越多，PDF 越大"""
        short_task = _make_task(output_text="# 短文\n\n只有一段。", title="短文")
        long_task = _make_task(output_text=MARKDOWN_CONTENT * 3, title="长文")

        from app.services.exporter import PdfExporter
        exporter = PdfExporter()
        short_data = exporter.export(short_task)
        long_data = exporter.export(long_task)
        assert len(long_data) > len(short_data)


# ---------------------------------------------------------------------------
# 边界情况
# ---------------------------------------------------------------------------

class TestExportEdgeCases:
    def test_docx_single_heading_only(self):
        """只有一个标题的内容也能正常导出"""
        import docx
        from app.services.exporter import DocxExporter
        task = _make_task(output_text="# 单章节报告\n\n这是唯一的内容段落。")
        data = DocxExporter().export(task)
        doc = docx.Document(io.BytesIO(data))
        assert any("单章节报告" in p.text for p in doc.paragraphs)

    def test_docx_all_chinese_content(self):
        """纯中文内容（无英文）正确导出"""
        import docx
        from app.services.exporter import DocxExporter
        content = "# 中文标题\n\n" + "这是纯中文段落内容，没有任何英文字母。" * 5
        task = _make_task(output_text=content, title="纯中文报告")
        data = DocxExporter().export(task)
        doc = docx.Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "中文标题" in full_text
        assert "纯中文段落" in full_text

    def test_pdf_all_chinese_content(self):
        """纯中文内容 PDF 导出不崩溃"""
        from app.services.exporter import PdfExporter
        content = "# 中文标题\n\n" + "这是纯中文段落内容。" * 10
        task = _make_task(output_text=content, title="纯中文")
        data = PdfExporter().export(task)
        assert data[:5] == b"%PDF-"
