"""Step 6.1 导出服务测试 — RED先跑失败，GREEN实现后全绿"""
from __future__ import annotations

import io
import uuid
from datetime import datetime, UTC
from unittest.mock import MagicMock

import pytest


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_task():
    task = MagicMock()
    task.id = uuid.uuid4()
    task.title = "量子计算技术报告"
    task.mode = "technical_report"
    task.status = "completed"
    task.output_text = """# 量子计算技术报告

## 第一章 概述

量子计算是一种利用量子力学原理的计算范式。

## 第二章 核心原理

量子比特（qubit）是量子计算的基本单元。

### 2.1 叠加态

量子叠加允许 qubit 同时处于 0 和 1 的状态。

```python
# 示例代码
print("hello quantum")
```

## 第三章 应用前景

量子计算在密码学、药物发现领域具有革命性潜力。
"""
    task.word_count = 100
    task.created_at = datetime(2026, 3, 26, 10, 0, 0)
    task.finished_at = datetime(2026, 3, 26, 10, 30, 0)
    return task


@pytest.fixture
def incomplete_task():
    task = MagicMock()
    task.id = uuid.uuid4()
    task.title = "未完成任务"
    task.status = "running"
    task.output_text = None
    return task


# ---------------------------------------------------------------------------
# BaseExporter
# ---------------------------------------------------------------------------

class TestBaseExporter:
    def test_base_exporter_is_abstract(self):
        from app.services.exporter import BaseExporter
        with pytest.raises(TypeError):
            BaseExporter()  # 抽象类不能直接实例化

    def test_base_exporter_has_export_method(self):
        from app.services.exporter import BaseExporter
        import inspect
        assert hasattr(BaseExporter, "export")
        assert inspect.isabstract(BaseExporter)


# ---------------------------------------------------------------------------
# DocxExporter
# ---------------------------------------------------------------------------

class TestDocxExporter:
    def test_export_returns_bytes(self, sample_task):
        from app.services.exporter import DocxExporter
        exporter = DocxExporter()
        result = exporter.export(sample_task)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_is_valid_docx(self, sample_task):
        """返回的 bytes 可以被 python-docx 正确加载"""
        from app.services.exporter import DocxExporter
        import docx
        exporter = DocxExporter()
        result = exporter.export(sample_task)
        doc = docx.Document(io.BytesIO(result))
        # 文档有段落内容
        assert len(doc.paragraphs) > 0

    def test_docx_contains_title(self, sample_task):
        from app.services.exporter import DocxExporter
        import docx
        exporter = DocxExporter()
        result = exporter.export(sample_task)
        doc = docx.Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert sample_task.title in all_text

    def test_docx_contains_headings(self, sample_task):
        from app.services.exporter import DocxExporter
        import docx
        exporter = DocxExporter()
        result = exporter.export(sample_task)
        doc = docx.Document(io.BytesIO(result))
        heading_texts = [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert len(heading_texts) >= 2

    def test_docx_raises_on_no_output_text(self, incomplete_task):
        from app.services.exporter import DocxExporter, ExportNotReadyError
        exporter = DocxExporter()
        with pytest.raises(ExportNotReadyError):
            exporter.export(incomplete_task)


# ---------------------------------------------------------------------------
# PdfExporter
# ---------------------------------------------------------------------------

class TestPdfExporter:
    def test_export_returns_bytes(self, sample_task):
        from app.services.exporter import PdfExporter
        exporter = PdfExporter()
        result = exporter.export(sample_task)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_is_valid_pdf(self, sample_task):
        """返回的 bytes 以 PDF 魔数开头"""
        from app.services.exporter import PdfExporter
        exporter = PdfExporter()
        result = exporter.export(sample_task)
        assert result[:4] == b"%PDF"

    def test_pdf_raises_on_no_output_text(self, incomplete_task):
        from app.services.exporter import PdfExporter, ExportNotReadyError
        exporter = PdfExporter()
        with pytest.raises(ExportNotReadyError):
            exporter.export(incomplete_task)

    def test_pdf_nonempty_output(self, sample_task):
        """PDF 内容超过最小阈值（非空白文档）"""
        from app.services.exporter import PdfExporter
        exporter = PdfExporter()
        result = exporter.export(sample_task)
        assert len(result) > 1024  # 至少 1KB


# ---------------------------------------------------------------------------
# MarkdownParser (内部工具)
# ---------------------------------------------------------------------------

class TestMarkdownParser:
    def test_parse_returns_blocks(self):
        from app.services.exporter import parse_markdown_blocks
        md = "# Title\n\nParagraph text.\n\n## Section\n\nMore text."
        blocks = parse_markdown_blocks(md)
        assert len(blocks) >= 3

    def test_parse_heading_block(self):
        from app.services.exporter import parse_markdown_blocks
        blocks = parse_markdown_blocks("# My Title")
        headings = [b for b in blocks if b["type"] == "heading"]
        assert len(headings) == 1
        assert headings[0]["level"] == 1
        assert headings[0]["text"] == "My Title"

    def test_parse_code_block(self):
        from app.services.exporter import parse_markdown_blocks
        md = "```python\nprint('hello')\n```"
        blocks = parse_markdown_blocks(md)
        code_blocks = [b for b in blocks if b["type"] == "code"]
        assert len(code_blocks) == 1

    def test_parse_paragraph(self):
        from app.services.exporter import parse_markdown_blocks
        blocks = parse_markdown_blocks("Just a paragraph.")
        paras = [b for b in blocks if b["type"] == "paragraph"]
        assert len(paras) == 1
        assert "Just a paragraph" in paras[0]["text"]

    def test_parse_empty_string(self):
        from app.services.exporter import parse_markdown_blocks
        blocks = parse_markdown_blocks("")
        assert blocks == []
