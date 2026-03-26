"""导出服务 — DOCX / PDF 导出器"""
from __future__ import annotations

import io
import re
from abc import ABC, abstractmethod
from typing import Any


class ExportNotReadyError(Exception):
    """任务尚未完成，无法导出"""


# ---------------------------------------------------------------------------
# Markdown 解析
# ---------------------------------------------------------------------------

def parse_markdown_blocks(text: str) -> list[dict[str, Any]]:
    """将 Markdown 文本解析为结构化块列表。

    每个块为 dict，包含 type 字段：
    - heading: level (int), text (str)
    - code:    language (str), text (str)
    - paragraph: text (str)
    """
    if not text or not text.strip():
        return []

    blocks: list[dict[str, Any]] = []
    # 先处理代码块（避免内部内容被误解析）
    code_pattern = re.compile(r"```(\w*)\n([\s\S]*?)```", re.MULTILINE)
    placeholders: dict[str, dict[str, Any]] = {}

    def replace_code(m: re.Match) -> str:
        key = f"__CODE_{len(placeholders)}__"
        placeholders[key] = {"type": "code", "language": m.group(1), "text": m.group(2).rstrip()}
        return key

    cleaned = code_pattern.sub(replace_code, text)

    for raw_line in cleaned.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        # 占位符还原
        if line in placeholders:
            blocks.append(placeholders[line])
            continue
        # 标题
        heading_m = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_m:
            blocks.append({"type": "heading", "level": len(heading_m.group(1)), "text": heading_m.group(2).strip()})
            continue
        # 普通段落
        blocks.append({"type": "paragraph", "text": line})

    return blocks


# ---------------------------------------------------------------------------
# 基类
# ---------------------------------------------------------------------------

class BaseExporter(ABC):
    @abstractmethod
    def export(self, task: Any) -> bytes:
        """将任务结果导出为字节流"""

    def _validate(self, task: Any) -> None:
        if not getattr(task, "output_text", None):
            raise ExportNotReadyError(f"Task {task.id} has no output_text (status={task.status})")


# ---------------------------------------------------------------------------
# DOCX 导出
# ---------------------------------------------------------------------------

class DocxExporter(BaseExporter):
    def export(self, task: Any) -> bytes:
        self._validate(task)
        import docx
        from docx.shared import Pt

        doc = docx.Document()

        # 标题页
        doc.add_heading(task.title, level=0)

        created = getattr(task, "created_at", None)
        if created:
            doc.add_paragraph(f"创建时间：{created.strftime('%Y-%m-%d %H:%M') if hasattr(created, 'strftime') else created}")

        finished = getattr(task, "finished_at", None)
        if finished:
            doc.add_paragraph(f"完成时间：{finished.strftime('%Y-%m-%d %H:%M') if hasattr(finished, 'strftime') else finished}")

        word_count = getattr(task, "word_count", 0)
        if word_count:
            doc.add_paragraph(f"字数：{word_count}")

        doc.add_page_break()

        # 正文内容
        for block in parse_markdown_blocks(task.output_text):
            btype = block["type"]
            if btype == "heading":
                level = min(block["level"], 9)
                doc.add_heading(block["text"], level=level)
            elif btype == "code":
                p = doc.add_paragraph(block["text"])
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            else:
                doc.add_paragraph(block["text"])

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF 导出
# ---------------------------------------------------------------------------

def _register_cjk_font() -> None:
    """注册中文 CID 字体（幂等，多次调用安全）"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


class PdfExporter(BaseExporter):
    def export(self, task: Any) -> bytes:
        self._validate(task)
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted

        _register_cjk_font()

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=25 * mm,
            rightMargin=25 * mm,
            topMargin=25 * mm,
            bottomMargin=25 * mm,
        )

        base_styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            "CustomNormal",
            parent=base_styles["Normal"],
            fontName="STSong-Light",
            fontSize=11,
            leading=18,
        )
        h1_style = ParagraphStyle(
            "CustomH1",
            parent=base_styles["Heading1"],
            fontName="STSong-Light",
            fontSize=18,
            leading=26,
            spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "CustomH2",
            parent=base_styles["Heading2"],
            fontName="STSong-Light",
            fontSize=15,
            leading=22,
            spaceAfter=4,
        )
        h3_style = ParagraphStyle(
            "CustomH3",
            parent=base_styles["Heading3"],
            fontName="STSong-Light",
            fontSize=13,
            leading=20,
            spaceAfter=3,
        )
        code_style = ParagraphStyle(
            "CustomCode",
            parent=base_styles["Code"],
            fontName="Courier",
            fontSize=9,
            leading=13,
        )

        heading_styles = {1: h1_style, 2: h2_style, 3: h3_style}

        story = []

        # 标题页
        story.append(Paragraph(task.title, h1_style))
        created = getattr(task, "created_at", None)
        if created:
            ts = created.strftime("%Y-%m-%d %H:%M") if hasattr(created, "strftime") else str(created)
            story.append(Paragraph(f"创建时间：{ts}", normal_style))
        story.append(Spacer(1, 10 * mm))

        # 正文
        for block in parse_markdown_blocks(task.output_text):
            btype = block["type"]
            if btype == "heading":
                level = block["level"]
                style = heading_styles.get(level, h3_style)
                story.append(Paragraph(block["text"], style))
            elif btype == "code":
                story.append(Preformatted(block["text"], code_style))
            else:
                # 转义 XML 特殊字符
                safe = block["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, normal_style))
            story.append(Spacer(1, 2 * mm))

        doc.build(story)
        return buf.getvalue()
